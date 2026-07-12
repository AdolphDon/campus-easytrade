"""
FastAPI接口：相当于接口层，获取信息-调用rag服务or知识库文档服务-返回信息
提供 REST API 供 Java 后端调用，实现：
  1. POST /api/chat         - 智能客服问答
  2. POST /api/knowledge/upload   - 上传文档到知识库
  3. DELETE /api/knowledge/delete - 从知识库删除文档

启动方式：uvicorn fastapi_service:app --host 0.0.0.0 --port 8000
"""
import os

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional

# 重要：将 streamlit 相关的导入放在最前面排除，避免冲突
# 此服务不依赖 Streamlit，所以不会有冲突

from rag import RagService  # 导入自己写的RAG知识库服务
from knowledge_base import KnowledgeBaseService, get_string_md5  # 导入知识库服务
import config_data as config  # 导入配置文件

# ========== FastAPI 应用初始化 ==========
app = FastAPI(title="智能客服AI服务", version="1.0.0",
              description="分流处理用户商品对话：带购买/数量语句由AI解析购物车参数；纯搜索语句交由前端渲染商品卡片")

# 跨域配置-允许Java后端调用-给FastAPI注册跨域中间件CORSMiddleware
app.add_middleware(
    CORSMiddleware, #FastAPI内置的跨域处理中间件，专门处理浏览器同源策略限制
    allow_origins=["*"], #允许所有来源域名访问接口
    allow_credentials=True, #允许请求携带凭证：Cookie、Token、Session、认证头信息
    allow_methods=["*"], #允许所有HTTP请求方式
    allow_headers=["*"], #允许前端携带任意请求头，比如Authorization存jwttoken、自定义业务请求头
)

# 声明两个全局变量，等待初始化时注入RAG检索增强服务、知识库管理服务到这个变量中
rag_service: Optional[RagService] = None #RAG检索增强服务
kb_service: Optional[KnowledgeBaseService] = None #知识库管理服务

# 给rag检索服务和知识库管理服务实例化(创建对象-方便调用)
@app.on_event("startup") #启动事件装饰器，服务启动的时候就执行一次
async def startup_event():
    global rag_service, kb_service #声明使用上边定义的全局变量
    rag_service = RagService() #启动前两个变量都是None，这里完成实例化，后续所有路由接口可直接全局调用
    kb_service = KnowledgeBaseService() #启动前两个变量都是None，这里完成实例化，后续所有路由接口可直接全局调用
    print("[FastAPI] 智能客服AI服务启动完成")
    print(f"[FastAPI] 聊天模型: {config.chat_model_name}")
    print(f"[FastAPI] 向量模型: {config.embedding_model_name}")
    print(f"[FastAPI] 向量数据库: {config.persist_directory}")

    # 启动自检：检查ChromaDB数据与md5记录是否一致
    try:
        chroma_count = len(kb_service.chroma.get().get("ids", []))
        md5_count = 0
        if os.path.exists(config.md5_path):
            with open(config.md5_path, "r", encoding="utf-8") as f:
                md5_count = len([l for l in f.readlines() if l.strip()])
        print(f"[FastAPI] 自检: ChromaDB 文档数={chroma_count}, md5 记录数={md5_count}")
        if md5_count > 0 and chroma_count == 0:
            print("[FastAPI] ⚠️ 警告: md5 有记录但 ChromaDB 为空，可能数据已丢失！")
            print("[FastAPI] ⚠️ 建议清空 md5.text 后重新上传文档")
    except Exception as e:
        print(f"[FastAPI] 自检异常: {e}")


# ========== 请求/响应DTO、VO模型 ==========
# BaseModel是Pydantic数据校验基类-接口入参自动类型校验
class ChatRequest(BaseModel):
    """面向用户聊天接口DTO规范"""
    input: str  #用户问题-用户原始输入文本
    session_id: str = "user_001" #会话唯一标识，用于区分不同用户
    order_context: str = "" #当前用户的未完成订单上下文（Java 端已预查好，直接传入）
    access_token: str = "" #JWT token（用于工具调用时识别用户身份）


class ChatResponse(BaseModel):
    """面向用户聊天接口VO规范"""
    answer: str  # AI 回答


class DeleteRequest(BaseModel):
    """删除文档接口DTO规范"""
    file_name: str  # 要删除的文件名


class DeleteResponse(BaseModel):
    """删除文档接口VO规范"""
    success: bool
    message: str #操作描述文案，给前端展示提示


class UploadResponse(BaseModel):
    """上传文档接口VO规范"""
    success: bool
    message: str #操作描述文案，给前端展示提示
    file_name: str = ""
    md5: str = ""
    content: str = ""  #提取的文本内容（用于预览）
    chunk_count: int = 0  #分块数量


# ========== API 端点 ==========
# 面向用户聊天
@app.post("/api/chat", response_model=ChatResponse)
# 这是智能客服核心对话POST接口，接收前端、Java传来的用户对话请求
# ChatResponse相当于上面自定义的VO，返回参数必须符合规范
async def chat(request: ChatRequest): #相当于自定义的DTO，请求参数需符合规范
    """
    智能客服问答
    接收用户问题 → RAG 检索知识库 → 调用大模型生成回答 → 返回结果
    """
    # ①校验rag服务是否实例化成功
    global rag_service #声明使用文件顶部全局变量，启动成功才实例化，失败就为None
    if rag_service is None:
        raise HTTPException(status_code=503, detail="AI服务未初始化")

    # ②校验java前端是否传进来了用户唯一标识id
    # 如果java前端传来的session_id有值就用真实值，空的就用"user_001"
    session_id = request.session_id or "user_001"

    # ③调用rag服务，传参并将返回值封装到answer中
    try:
        # 使用Tool Calling版本的问答（RAG检索 + 按需调Java接口查实时数据）
        answer = rag_service.chat_with_tools( #调用rag服务chat_with_tools方法
            input_text=request.input, #用户原始提问
            session_id=session_id, #用户会话标识，维护上下文记忆
            order_context=request.order_context, #Java后端提前查好的用户订单文本，直接塞进Prompt，减少远程重复调用
            access_token=request.access_token, #用户JWT凭证
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI服务调用失败: {str(e)}")

    return ChatResponse(answer=answer) #按要求返回数据

# 知识库文档管理
@app.post("/api/knowledge/upload", response_model=UploadResponse)
# 知识库文档上传接口
async def upload_knowledge(
        file: UploadFile = File(...), #FastAPI内置文件上传专用类型，封装文件名、文件二进制内容、后缀、大小
        operator: str = Form("管理员"), #默认值，前端不传operator时，操作人自动填充为「管理员」
):
    """
    上传文档到知识库
    支持 TXT / PDF / DOCX / MD → 提取文本 → 分块 → 向量化 → 存入 ChromaDB
    """
    # ①校验知识库管理服务是否实例化成功
    global kb_service #声明使用文件顶部全局变量，启动成功才实例化，失败就为None
    if kb_service is None:
        raise HTTPException(status_code=503, detail="知识库服务未初始化")

    # ②检查文件类型
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")
    lower_name = file.filename.lower()
    allowed = [".txt", ".pdf", ".docx", ".md"]
    if not any(lower_name.endswith(ext) for ext in allowed):
        raise HTTPException(status_code=400, detail=f"不支持的文件格式，仅支持: {', '.join(allowed)}")

    # ③读取原始字节
    try:
        content_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"文件读取失败: {str(e)}")

    # ④检查文件内容是否为空
    if not content_bytes:
        raise HTTPException(status_code=400, detail="文件内容为空")

    # 根据文件类型提取文本
    content = ""
    try:
        if lower_name.endswith(".txt") or lower_name.endswith(".md"):
            content = content_bytes.decode("utf-8")
        elif lower_name.endswith(".pdf"):
            try:
                import fitz  # PyMuPDF
                import io
                pdf_doc = fitz.open(stream=content_bytes, filetype="pdf")
                pages = []
                for page in pdf_doc:
                    pages.append(page.get_text())
                content = "\n\n".join(pages)
                pdf_doc.close()
                if not content.strip():
                    raise HTTPException(status_code=400, detail="PDF 文件内容为空（无可提取的文本）")
            except ImportError:
                raise HTTPException(status_code=500, detail="PDF 解析库未安装，请执行: pip install PyMuPDF")
        elif lower_name.endswith(".docx"):
            try:
                import docx
                import io
                doc = docx.Document(io.BytesIO(content_bytes))
                paragraphs = [p.text for p in doc.paragraphs]
                content = "\n".join(paragraphs)
                if not content.strip():
                    # 尝试读取表格内容
                    tables = []
                    for table in doc.tables:
                        for row in table.rows:
                            cells = [cell.text for cell in row.cells]
                            tables.append(" | ".join(cells))
                    content = "\n".join(tables)
                if not content.strip():
                    raise HTTPException(status_code=400, detail="DOCX 文件内容为空（无可提取的文本）")
            except ImportError:
                raise HTTPException(status_code=500, detail="DOCX 解析库未安装，请执行: pip install python-docx")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"文件解析失败: {str(e)}")

    # 计算 MD5
    md5_hex = get_string_md5(content)

    # 上传到向量知识库
    try:
        result = kb_service.upload_by_str(content, file.filename, operator)
        # 解析结果获取分块数
        success = "[成功]" in result
        chunk_count = 0
        if success:
            # 从 ChromaDB 查询该文档的分块数
            try:
                query_result = kb_service.chroma.get(where={"source": file.filename})
                chunk_count = len(query_result.get("ids", []))
            except:
                pass
        return UploadResponse(
            success=success,
            message=result,
            file_name=file.filename,
            md5=md5_hex,
            content=content[:5000],  # 只存前5000字符用于预览
            chunk_count=chunk_count,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"知识库上传失败: {str(e)}")

# 知识库文档管理
@app.delete("/api/knowledge/delete", response_model=DeleteResponse)
async def delete_knowledge(request: DeleteRequest):
    """
    从知识库删除文档
    根据文件名删除 ChromaDB 中对应的所有向量片段
    """
    # ①校验知识库管理服务是否实例化成功
    global kb_service #声明使用文件顶部全局变量，启动成功才实例化，失败就为None
    if kb_service is None:
        raise HTTPException(status_code=503, detail="知识库服务未初始化")

    try:
        # 通过 ChromaDB 的 metadata 过滤获取要删除的文档ID
        # metadata 中的 source 字段存储了文件名
        result = kb_service.chroma.get(where={"source": request.file_name})
        ids = result.get("ids", [])

        if not ids:
            return DeleteResponse(
                success=False,
                message=f"未找到文档 '{request.file_name}' 的相关记录",
            )

        # 获取被删除文件的 MD5（从被删片段的 metadata 中读取）
        deleted_metadatas = result.get("metadatas", [])
        md5_to_clean = set()
        for m in deleted_metadatas:
            if m and m.get("md5"):
                md5_to_clean.add(m["md5"])

        # 执行删除
        kb_service.chroma.delete(ids=ids)

        # 检查每个 MD5 是否还被其他文档引用
        for md5_val in md5_to_clean:
            remaining = kb_service.chroma.get(where={"md5": md5_val})
            if not remaining.get("ids"):
                # 没有其他文档引用此 MD5 → 从 md5.text 移除
                md5_path = config.md5_path
                if os.path.exists(md5_path):
                    with open(md5_path, "r", encoding="utf-8") as f:
                        lines = f.readlines()
                    new_lines = [l for l in lines if l.strip() != md5_val]
                    with open(md5_path, "w", encoding="utf-8") as f:
                        f.writelines(new_lines)
                    print(f"[FastAPI] MD5 {md5_val} 已从 md5.text 移除，无其他文档引用")

        return DeleteResponse(
            success=True,
            message=f"成功删除文档 '{request.file_name}'，共移除 {len(ids)} 个向量片段",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"知识库删除失败: {str(e)}")


@app.get("/api/health")
async def health_check():
    """健康检查接口：开发、后端联调排查 运维、服务器容器自检"""
    return {
        "status": "ok",
        "chat_model": config.chat_model_name,
        "embedding_model": config.embedding_model_name,
        "chroma_db": config.persist_directory,
    }


# ========== 直接启动入口 ==========
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
